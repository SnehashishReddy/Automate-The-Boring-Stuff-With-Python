import docx

file = open('guests.txt','r')
lst = list(file.readlines())
doc = docx.Document()
for x in lst:
    y = list(x)
    if '\n' in y:
        y.pop(-1)
    sol1 = doc.add_paragraph("","Normal")
    sol1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    sol1.add_run("It would be a pleasure to have the company of").italic = True
    sol1 = doc.add_paragraph("","Normal") #
    sol1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    sol2 = sol1.add_run(y)
    sol2.size = docx.shared.Pt(30)
    sol2.bold = True
    sol2.all_caps = True
    sol1 = doc.add_paragraph("","Normal")
    sol1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    sol1.add_run("April 1st")
    sol1 = doc.add_paragraph("","Normal")
    sol1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    sol1.add_run("at 7 o' clock").italic = True
    doc.add_page_break()
    doc.save("Invitations.docx")
file.close()
